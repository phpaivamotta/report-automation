from docx import Document
from docx.shared import Pt

def add_bullets_above_tables(docx_path, output_path):
    # Open the existing document
    doc = Document(docx_path)

    # Find all the tables in the document
    for table in doc.tables:
        # Get the index of the table in the list of elements
        table_idx = find_table_index(doc, table)

        # Insert two bullet points before the table
        bullet_1 = doc.paragraphs[table_idx].insert_paragraph_before("• Bullet point 1")
        apply_bullet_style(bullet_1)

        bullet_2 = bullet_1.insert_paragraph_before("• Bullet point 2")
        apply_bullet_style(bullet_2)

    # Save the modified document
    doc.save(output_path)
    print(f"Bullet points added above each table and saved to {output_path}")

def find_table_index(doc, table):
    """
    Finds the index of the table in the document to determine where to insert the bullet points.
    """
    for i, paragraph in enumerate(doc.paragraphs):
        if table._element.getprevious() == paragraph._element:
            return i
    return -1

def apply_bullet_style(paragraph):
    """
    Apply bullet-style formatting to a paragraph with Calibri font and size 12.
    """
    run = paragraph.runs[0]
    run.font.name = 'Calibri (Body)'
    run.font.size = Pt(12)

# Example usage
docx_file_path = r"C:\Users\phpai\OneDrive\Desktop\report-automation\modified_document7.docx"
output_docx_path = r"C:\Users\phpai\OneDrive\Desktop\report-automation\modified_document_with_bullets.docx"

add_bullets_above_tables(docx_file_path, output_docx_path)
