import os
import csv
import logging
from docx import Document
from docx.oxml.ns import qn
from collections import OrderedDict
from dotenv import load_dotenv

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def get_next_report_id(report_csv_path):
    """Determines the next available report ID."""
    try:
        if not os.path.exists(report_csv_path):
            return 1
        
        with open(report_csv_path, 'r', newline='', encoding='utf-8') as report_csv:
            reader = csv.DictReader(report_csv)
            report_ids = [int(row['Report ID']) for row in reader if row['Report ID'].isdigit()]
            return max(report_ids) + 1 if report_ids else 1
    except Exception as e:
        logging.error(f"Error determining next report ID: {str(e)}")
        raise

def create_report_folder(base_output_folder, report_id):
    """Creates a folder for the current report using the report ID."""
    try:
        report_folder = os.path.join(base_output_folder, f"report_{report_id:04d}")
        os.makedirs(report_folder, exist_ok=True)
        return report_folder
    except Exception as e:
        logging.error(f"Error creating report folder: {str(e)}")
        raise

def extract_report_data(docx_path, report_folder):
    """Extracts report data and images from the DOCX template."""
    try:
        document = Document(docx_path)
        report_data = {
            'Customer': '', 'Subject': '', 'Customer Address': '', 'From': '',
            'Customer Contact': '', 'Company': '', 'Inspection Site': '',
            'Maverick Contact Info': '', 'Customer PO No.': '', 'Maverick Job': '',
            'Customer CCs': '', 'Maverick CCs': '', 'Inspection Date(s)': '',
            'Report Date': '', 'Introduction': '', 'Entrance Meeting': '',
            'Drawings Used': OrderedDict(), 'Specifications Used': OrderedDict(),
            'Conclusions': '', 'Pictures': []
        }

        current_section = None
        
        for table in document.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) == 4:
                    current_section = process_four_column_row(cells, report_data, current_section)
                elif len(cells) == 3:
                    process_picture_row(cells, report_data, report_folder)

        # Convert OrderedDicts back to lists
        report_data['Drawings Used'] = list(report_data['Drawings Used'].keys())
        report_data['Specifications Used'] = list(report_data['Specifications Used'].keys())

        return report_data
    except Exception as e:
        logging.error(f"Error extracting report data: {str(e)}")
        raise

def process_four_column_row(cells, report_data, current_section):
    """Process a row with four columns."""
    left_descriptor = cells[0].text.strip().rstrip(':')
    left_value = cells[1].text.strip()
    right_descriptor = cells[2].text.strip().rstrip(':')
    right_value = cells[3].text.strip()

    for descriptor, value in [(left_descriptor, left_value), (right_descriptor, right_value)]:
        if descriptor in report_data and isinstance(report_data[descriptor], str):
            report_data[descriptor] = value

    if "Drawing" in left_descriptor or "Drawing" in right_descriptor:
        current_section = 'Drawings Used'
    elif "Specification" in left_descriptor or "Specification" in right_descriptor:
        current_section = 'Specifications Used'
    elif "Conclusions" in left_descriptor or "Conclusions" in right_descriptor:
        current_section = 'Conclusions'
        report_data['Conclusions'] = left_value if "Conclusions" in left_descriptor else right_value
    
    if current_section == 'Drawings Used':
        for value in [left_value, right_value]:
            if value and "Drawing" not in value:
                report_data[current_section][value] = None
    elif current_section == 'Specifications Used':
        for value in [left_value, right_value]:
            if value and "Specification" not in value:
                report_data[current_section][value] = None

    return current_section

def process_picture_row(cells, report_data, report_folder):
    """Process a row containing picture information."""
    if len(cells) == 3 and cells[0].text.strip() == 'Description' and cells[1].text.strip() == 'Caption' and cells[2].text.strip() == 'Picture':
        return  # This is the header row, skip it

    description = cells[0].text.strip()
    caption = cells[1].text.strip()
    picture_cell = cells[2]

    image_name = extract_and_save_image(picture_cell, report_folder, len(report_data['Pictures']) + 1)
    if image_name:
        report_data['Pictures'].append({
            'Description': description,
            'Caption': caption,
            'Picture File Name': image_name
        })

def extract_and_save_image(picture_cell, report_folder, image_count):
    """Extract and save an image from a cell."""
    for paragraph in picture_cell.paragraphs:
        for run in paragraph.runs:
            if 'graphicData' in run._element.xml:
                blip = next((elem for elem in run._element.iter() if elem.tag.endswith('blip')), None)
                if blip is not None:
                    try:
                        embed = blip.get(qn('r:embed'))
                        image_part = run.part.related_parts[embed]
                        image_extension = os.path.splitext(image_part.partname)[1]
                        image_name = f"image_{image_count:03d}{image_extension}"
                        image_path = os.path.join(report_folder, image_name)
                        
                        with open(image_path, 'wb') as f:
                            f.write(image_part.blob)
                        
                        logging.info(f"Saved image: {image_name}")
                        return image_name
                    except Exception as e:
                        logging.error(f"Error saving image: {str(e)}")
    return None

def append_to_csv(report_data, report_id, report_csv_path, picture_csv_path, report_folder):
    """Appends report and picture data to the corresponding CSV files."""
    try:
        # Append report data
        with open(report_csv_path, 'a', newline='', encoding='utf-8') as report_csv:
            writer = csv.writer(report_csv)
            if os.path.getsize(report_csv_path) == 0:  # Write header if the file is empty
                writer.writerow(['Report ID', 'Customer', 'Subject', 'Customer Address', 'From', 
                                 'Customer Contact', 'Company', 'Inspection Site', 'Maverick Contact Info', 
                                 'Customer PO No.', 'Maverick Job', 'Customer CCs', 'Maverick CCs', 
                                 'Inspection Date(s)', 'Report Date', 'Introduction', 'Entrance Meeting', 
                                 'Drawings Used', 'Specifications Used', 'Conclusions', 'Report Folder'])
            
            drawings = '; '.join(report_data['Drawings Used'])
            specifications = '; '.join(report_data['Specifications Used'])
            
            writer.writerow([report_id, report_data['Customer'], report_data['Subject'], 
                             report_data['Customer Address'], report_data['From'], 
                             report_data['Customer Contact'], report_data['Company'], 
                             report_data['Inspection Site'], report_data['Maverick Contact Info'], 
                             report_data['Customer PO No.'], report_data['Maverick Job'], 
                             report_data['Customer CCs'], report_data['Maverick CCs'], 
                             report_data['Inspection Date(s)'], report_data['Report Date'], 
                             report_data['Introduction'], report_data['Entrance Meeting'], 
                             drawings, specifications, report_data['Conclusions'],
                             os.path.basename(report_folder)])

        # Append picture data
        with open(picture_csv_path, 'a', newline='', encoding='utf-8') as picture_csv:
            writer = csv.writer(picture_csv)
            if os.path.getsize(picture_csv_path) == 0:  # Write header if the file is empty
                writer.writerow(['Picture ID', 'Report ID', 'Description', 'Caption', 'Picture File Name', 'Report Folder'])
            
            for index, picture in enumerate(report_data['Pictures'], start=1):
                writer.writerow([f"{report_id}_{index}", report_id, picture['Description'], 
                                 picture['Caption'], picture['Picture File Name'], os.path.basename(report_folder)])

        logging.info(f"Data successfully written to CSV files.")
    except Exception as e:
        logging.error(f"Error writing to CSV files: {str(e)}")
        raise

def main():
    # Use environment variables
    load_dotenv(override=True)
    base_output_folder = os.getenv('EXTRACTED_DATA_PATH')
    docx_path = os.getenv('INPUT_DOC_PATH')
    report_csv_path = os.path.join(base_output_folder, 'report_info.csv')
    picture_csv_path = os.path.join(base_output_folder, 'picture_info.csv')

    # Check if input file exists
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Input DOCX file not found: {docx_path}")

    # Get the next report ID
    report_id = get_next_report_id(report_csv_path)

    # Create a folder for this report using the report ID
    report_folder = create_report_folder(base_output_folder, report_id)

    # Extract and append data
    report_data = extract_report_data(docx_path, report_folder)
    append_to_csv(report_data, report_id, report_csv_path, picture_csv_path, report_folder)

if __name__ == "__main__":
    main()