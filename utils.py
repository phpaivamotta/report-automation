from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import win32com.client as win32
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
import time
import os
import glob



def add_table_with_images(doc, header_text, table_counter, num_cols, image_path1, image_path2=None):

    if table_counter == 0:
        # Find the paragraph with the header text
        target_paragraph = None
        for paragraph in doc.paragraphs:
            if header_text in paragraph.text:
                target_paragraph = paragraph
                break
    else:
        # Find the last table in the document
        last_table = doc.tables[-1]  # Get the last table

        # Get the last table's XML element
        tbl_element = last_table._element

        # Create a new paragraph XML element
        new_paragraph_element = OxmlElement('w:p')

        # Insert the new paragraph right after the last table
        tbl_element.addnext(new_paragraph_element)

        # Create a paragraph object that we can use programmatically for further insertion
        target_paragraph = doc.add_paragraph()
        target_paragraph._element = new_paragraph_element


    if target_paragraph is None:
        print(f"Header '{header_text}' not found in the document.")
        return

    target_paragraph.insert_paragraph_before()

    # Add a table after the new paragraph
    # We can't directly control placement through doc.add_table, so we'll insert it programmatically
    table = doc.add_table(rows=1, cols=num_cols)

    set_table_borders(table)

     # Disable automatic table resizing
    table.autofit = False

    # Set table alignment to center
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for column in table.columns:
        column.width = Inches(3.7)  # Set column width to 2 inches
        for cell in column.cells:
            cell.width = Inches(3.7)

    # Move the table to the specific location after the target paragraph
    # Using the XML elements for moving the table
    target_paragraph._element.addnext(table._element)

    # Add images to each cell in the table
    if num_cols == 1:
        cell = table.cell(0, 0)
        set_cell_margins(table, left=72, right=72, top=72, bottom=0)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path1, width=Inches(3.6))  # Adjust width as needed

    elif num_cols == 2:
        cell = table.cell(0, 0)
        set_cell_margins(table, left=72, right=72, top=72, bottom=0)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path1, width=Inches(3.6))  # Adjust width as needed

        cell = table.cell(0, 1)
        set_cell_margins(table, left=72, right=72, top=72, bottom=0)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path2, width=Inches(3.6))  # Adjust width as needed

    print("Table with images were added successfully.")


def replace_text_in_paragraph(paragraph, old_texts, new_texts):
    for old_text, new_text in zip(old_texts, new_texts):
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

            run = paragraph.runs[0]
            run.font.name = 'Calibri (Body)'
            run.font.size = Pt(11)


def replace_text_in_table(table, old_texts, new_texts):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, old_texts, new_texts)

    print("Project details in Table 1 were modified successfully.")


def add_captions_with_win32com(doc_path, i, num_cols, image_path1, image_path2=None):
    # Open Word application
    word = win32.Dispatch('Word.Application')
    word.Visible = True  # Set to True if you want to see Word while working

    # Open the existing document
    doc = word.Documents.Open(doc_path)

    # Loop through all inline shapes (images) in the document
    # for i, inline_shape in enumerate(doc.InlineShapes):
    
    if num_cols == 1:
        # Get image file name
        file_name = os.path.basename(image_path1)
        file_name = ". " + file_name
        # Select the inline shape (image)
        doc.InlineShapes(i+1).Select() # The InlineShape() method is 1-based indexed
        # Insert a caption for the selected image (exclude custom title, only label + number)
        word.Selection.InsertCaption(Label="Figure", Title=file_name, ExcludeLabel=False, Position=-1)
        # Move the selection to the end of the caption
        word.Selection.MoveRight(Unit=2, Count=1, Extend=1)
        # Delete any text after the caption label (if any text remains after "Figure X")
        word.Selection.TypeBackspace()

    elif num_cols ==  2:
        # Get image file name
        file_name = os.path.basename(image_path1)
        file_name = ". " + file_name
        # Select the inline shape (image)
        doc.InlineShapes(i+1).Select() # The InlineShape() method is 1-based indexed
        # Insert a caption for the selected image (exclude custom title, only label + number)
        word.Selection.InsertCaption(Label="Figure", Title=file_name, ExcludeLabel=False, Position=-1)
        # Move the selection to the end of the caption
        word.Selection.MoveRight(Unit=2, Count=1, Extend=1)
        # Delete any text after the caption label (if any text remains after "Figure X")
        word.Selection.TypeBackspace()

        # Get image file name
        file_name = os.path.basename(image_path2)
        file_name = ". " + file_name
        # Select the inline shape (image)
        doc.InlineShapes(i+2).Select() # The InlineShape() method is 1-based indexed
        # Insert a caption for the selected image (exclude custom title, only label + number)
        word.Selection.InsertCaption(Label="Figure", Title=file_name, ExcludeLabel=False, Position=-1)
        # Move the selection to the end of the caption
        word.Selection.MoveRight(Unit=2, Count=1, Extend=1)
        # Delete any text after the caption label (if any text remains after "Figure X")
        word.Selection.TypeBackspace()

    # Update all fields (important for cross-references)
    doc.Fields.Update()

    # Save and close the document
    doc.Save()
    doc.Close()
    word.Quit()

    print("Captions added successfully.")


def set_table_borders(table):
    """
    Set the borders of the table to the specified color with a width of 2.25 pt.
    """
    # Calculate the border size in eighths of a point
    border_size = 18  # 2.25 pt * 8 = 18 units

    # Define the border color
    border_color = "002060"  # Hex color code for RGB(0,32,96)

    # Define the border style XML
    tbl_borders = parse_xml(r'''
        <w:tblBorders %s>
            <w:top w:val="single" w:sz="%d" w:space="0" w:color="%s"/>
            <w:left w:val="single" w:sz="%d" w:space="0" w:color="%s"/>
            <w:bottom w:val="single" w:sz="%d" w:space="0" w:color="%s"/>
            <w:right w:val="single" w:sz="%d" w:space="0" w:color="%s"/>
            <w:insideH w:val="single" w:sz="%d" w:space="0" w:color="%s"/>
            <w:insideV w:val="single" w:sz="%d" w:space="0" w:color="%s"/>
        </w:tblBorders>
        ''' % (
            nsdecls('w'),
            border_size, border_color,
            border_size, border_color,
            border_size, border_color,
            border_size, border_color,
            border_size, border_color,
            border_size, border_color
        ))

    # Access or create the table properties element
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)

    # Remove any existing borders element
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)

    # Append the borders element to the table properties
    tblPr.append(tbl_borders)


def set_cell_margins(table, left=0, right=0, top=0, bottom=0):
    tc = table._element
    tblPr = tc.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')
    kwargs = {"left":left, "right":right, "top":top, "bottom":bottom}
    for m in ["left","right", "top", "bottom"]:
        node = OxmlElement("w:{}".format(m))
        node.set(qn('w:w'), str(kwargs.get(m)))
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)

    tblPr.append(tblCellMar)


def add_bullets_above_tables(output_doc_file_path, table_counter, num_cols):

    doc = Document(output_doc_file_path)

    # Loop through all tables in the document
    table = doc.tables[table_counter+1] # Table 1 should be skipped since it is the properties table

    # Find the paragraph just before the table
    paragraph_before_table = table._element.getprevious()

    if paragraph_before_table is not None:
        if num_cols == 2:
            # Insert two bullet points above the table
            bullet_1 = doc.add_paragraph("Bullet point 1", style='List Bullet 2')
            bullet_2 = doc.add_paragraph("Bullet point 2", style='List Bullet 2')
            
            # Insert the bullet points before the table
            paragraph_before_table.addnext(bullet_2._element)
            bullet_2._element.addprevious(bullet_1._element)

        if num_cols == 1:
            # Insert two bullet points above the table
            bullet_1 = doc.add_paragraph("Bullet point 1", style='List Bullet 2')
            
            # Insert the bullet points before the table
            paragraph_before_table.addnext(bullet_1._element)

    # doc.save(output_doc_file_path)
    print(f"Bullets added above all tables except the first")
    doc.save(output_doc_file_path)
    return doc

    # doc.save(output_doc_file_path)


def append_cross_references_to_bullets(docx_path, i):
    """Append cross-references to the beginning of each bullet point and make Figure 1 and Figure 2 bold."""
    # Open Word application
    word = win32.Dispatch('Word.Application')
    word.Visible = True  # Set to True if you want to see Word while working

    # Open the existing document
    doc = word.Documents.Open(docx_path)

    # Set the figure references for bullet 1 and bullet 2
    figure_1_ref = i + 1  # Reference to Figure 1
    figure_2_ref = i + 2  # Reference to Figure 2

    # Loop through the paragraphs to find bullet points and append cross-references
    for para in doc.Paragraphs:
        if para.Range.Text.strip() == "Bullet point 1":
            # Move the cursor to the beginning of the paragraph and insert the cross-reference for Figure 1
            word.Selection.SetRange(para.Range.Start, para.Range.Start)
            word.Selection.InsertCrossReference(
                ReferenceType="Figure",
                ReferenceKind=3,  # 3 corresponds to wdOnlyLabelAndNumber (Figure X)
                ReferenceItem=figure_1_ref,
                InsertAsHyperlink=True,  # Optional: make it a hyperlink
                IncludePosition=False,
                SeparateNumbers=False,
                SeparatorString=" "
            )

            # Bold the inserted Figure 1 text
            word.Selection.MoveLeft(Unit=1, Count=1, Extend=True)  # 1 = wdCharacter
            word.Selection.Font.Bold = True

            # Move cursor to the end of the bolded Figure 1 text
            word.Selection.Collapse(Direction=0)

            word.Selection.TypeText(" ")  # Add a space before the original text

            word.Selection.Font.Bold = False

            word.Selection.TypeText("shows ")  # Add a space before the original text

            # Move the selection to the end of the caption
            word.Selection.MoveRight(Unit=1, Count=14, Extend=1) #Bullet point 1

            # Delete any text after the caption label (if any text remains after "Figure X")
            word.Selection.TypeBackspace()

            set_font_formatting(para, word)

            set_paragraph_spacing(para, word)

        elif para.Range.Text.strip() == "Bullet point 2":
            # Move the cursor to the beginning of the paragraph and insert the cross-reference for Figure 2
            word.Selection.SetRange(para.Range.Start, para.Range.Start)
            word.Selection.InsertCrossReference(
                ReferenceType="Figure",
                ReferenceKind=3,  # 3 corresponds to wdOnlyLabelAndNumber (Figure X)
                ReferenceItem=figure_2_ref,
                InsertAsHyperlink=True,  # Optional: make it a hyperlink
                IncludePosition=False,
                SeparateNumbers=False,
                SeparatorString=" "
            )

            # Bold the inserted Figure 2 text
            word.Selection.MoveLeft(Unit=1, Count=1, Extend=True)  # 1 = wdCharacter
            word.Selection.Font.Bold = True

            # Move cursor to the end of the bolded Figure 2 text
            word.Selection.Collapse(Direction=0)

            word.Selection.TypeText(" ")  # Add a space before the original text

            word.Selection.Font.Bold = False

            word.Selection.TypeText("shows ")  # Add a space before the original text

            # Move the selection to the end of the caption
            word.Selection.MoveRight(Unit=1, Count=14, Extend=1) #Bullet point 1

            # Delete any text after the caption label (if any text remains after "Figure X")
            word.Selection.TypeBackspace()

            set_font_formatting(para, word)

            set_paragraph_spacing(para, word)

    # Save the document with cross-references
    doc.SaveAs(docx_path)
    doc.Close()
    word.Quit()

    print(f"Cross-references added.")


def set_font_formatting(para, word):
    """Set font formatting for the paragraph to Calibri 12."""
    # Apply the font to the whole range of the paragraph
    para.Range.Font.Name = 'Calibri (Body)'
    para.Range.Font.Size = 12

def set_paragraph_spacing(para, word):
    word.Selection.SetRange(para.Range.Start, para.Range.End)
    word.Selection.ParagraphFormat.SpaceBefore = 6
    word.Selection.ParagraphFormat.SpaceAfter = 6
    para.Style.NoSpaceBetweenParagraphsOfSameStyle = False


def delete_template_bullets(output_doc_file_path):

    doc = Document(output_doc_file_path)

    count = 0
    for para in doc.paragraphs:
        if para.style.name in ["List Bullet", "List Bullet 2", "List Bullet 3"]:
            if count >= 3:
                break
            # Remove the paragraph from the parent element (body)
            p = para._element
            p.getparent().remove(p)
            # Clean up after removing
            p._element = p = None
            count += 1

    doc.save(output_doc_file_path)
    print("Template bullets deleted.")
    


def get_images_from_folder(folder_path):
    # Define the image extensions you want to search for
    image_extensions = ['*.jpg', '*.jpeg', '*.png', '*.gif', '*.bmp', '*.tiff']

    # List to store the image file paths
    image_paths = []
    # Loop through each image extension to grab matching files
    for extension in image_extensions:
        image_paths.extend(glob.glob(os.path.join(folder_path, extension)))

    image_paths = sorted(image_paths, key=lambda x: x.lower())

    return image_paths

def delete_paragraph(paragraph):
    # Access the XML element of the paragraph
    p = paragraph._element
    # Access the parent element of the paragraph (usually the document body)
    p.getparent().remove(p)
    # Clean up after removing
    paragraph._element = None


def remove_empty_paragraphs_after_table(output_doc_file_path):
    """
    This function checks for empty paragraphs after tables and removes them.
    """

    doc = Document(output_doc_file_path)

    for i, table in enumerate(doc.tables):
        if i == 0:
            continue
        
        # Get the next element after the table
        next_element = table._element.getnext()

        # Continue checking for next paragraphs as long as they exist and are paragraphs
        while next_element is not None and next_element.tag.endswith('p'):
            # Check if the next element is a paragraph and is effectively empty (strip spaces and non-breaking spaces)
            paragraph_text = "".join(next_element.itertext()).strip()
            if not paragraph_text:
                # If it's an empty paragraph, remove it
                parent = next_element.getparent()
                parent.remove(next_element)

                # Get the next element after the removed one
                next_element = table._element.getnext()
            else:
                break  # Exit the loop if the paragraph is not empty

    # Save the document after making modifications
    doc.save(output_doc_file_path)
    # return doc


def remove_first_empty_paragraph_above_text(output_doc_file_path, text):
    """
    Removes the first empty paragraph above the given text in the document.
    """

    doc = Document(output_doc_file_path)

    # Iterate through paragraphs to find the one containing the target text
    for i, paragraph in enumerate(doc.paragraphs):
        if text in paragraph.text:
            # Check for the preceding paragraph
            prev_paragraph = paragraph._element.getprevious()

            # If the previous element is an empty paragraph, remove it
            if prev_paragraph is not None and prev_paragraph.tag.endswith('p'):
                prev_text = "".join(prev_paragraph.itertext()).strip()
                if not prev_text:
                    # Remove the first empty paragraph found
                    parent = prev_paragraph.getparent()
                    parent.remove(prev_paragraph)
                break  # Stop after deleting the first empty paragraph
            break  # Exit once the target paragraph is found

    doc.save(output_doc_file_path)
    # return doc
