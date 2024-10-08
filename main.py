from docx import Document
from docx.opc.coreprops import CoreProperties
from docx.shared import Inches
from utils import add_table_with_images
from utils import replace_text_in_table
from utils import add_captions_with_win32com
from utils import add_bullets_above_tables
from utils import append_cross_references_to_bullets
from utils import delete_template_bullets
from dotenv import load_dotenv
import os



# Load environment variables from .env file
load_dotenv(override=True)
template_file_path = os.getenv('TEMPLATE_DOC_PATH')
output_doc_file_path = os.getenv('OUTPUT_REPORT_DOC_PATH')

# Inputs
image_path_1 = r"C:\Users\phpai\OneDrive\Desktop\report-automation\Images\image1.jpeg"
image_path_2 = r"C:\Users\phpai\OneDrive\Desktop\report-automation\Images\image2.jpeg"

images = [
    image_path_1,
    image_path_2
]

doc_core_properties = {
    "title": "Title", # customer
    "author": "author", # from
    "subject": "subject", #
    "keywords": "keywords", # maverick job num
}

custom_properties = {
    "customer address": "customer address",
    "customer contact": "customer contact",
    "inspection site": "inspection site",
    "customer po num": "customer po num",
    "customer ccs": "customer css",
    "inspection date": "inspection date",
    "company": "company",
    "maverick contact info": "maverick contact info",
    "maverick ccs": "maverick ccs",
    "report date": "report date",
}

if __name__ == "__main__":

    # Open an existing document
    doc = Document(template_file_path)

    # Modify core properties
    core_properties = doc.core_properties
    # Set core properties
    core_properties.title = doc_core_properties["title"]
    core_properties.author = doc_core_properties["author"]
    core_properties.subject = doc_core_properties["subject"]
    core_properties.keywords = doc_core_properties["keywords"]

    # Iterate through all tables (if any)
    for table in doc.tables:
        replace_text_in_table(table, custom_properties.keys(), custom_properties.values())

    add_table_with_images(doc, "Inspection Observations:", image_path_1, image_path_2)

    # Save the modified document
    doc.save(output_doc_file_path)

    # Add captions to images
    add_captions_with_win32com(output_doc_file_path, images)

    # Add bullets above table
    doc = add_bullets_above_tables(output_doc_file_path)

    delete_template_bullets(doc, output_doc_file_path)

    # Add cross references
    append_cross_references_to_bullets(output_doc_file_path)